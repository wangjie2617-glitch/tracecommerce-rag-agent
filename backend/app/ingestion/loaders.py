"""Local PDF, text, Markdown, HTML, and DOCX loaders."""

from __future__ import annotations

import asyncio
import io
from datetime import UTC, datetime
from pathlib import PurePath

from bs4 import BeautifulSoup

from app.core.exceptions import AppError
from app.rag.types import ParsedDocument

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".html", ".htm", ".docx"}


def _load_sync(filename: str, content: bytes) -> str:
    extension = PurePath(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise AppError("不支持的文件类型", code="unsupported_file_type", status_code=415)
    if extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"## Page {index}\n\n{text}")
        return "\n\n".join(pages)
    if extension == ".docx":
        from docx import Document as DocxDocument

        document = DocxDocument(io.BytesIO(content))
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name.lower() if paragraph.style else ""
            lines.append(f"## {text}" if "heading" in style else text)
        return "\n\n".join(lines)
    decoded = content.decode("utf-8", errors="replace")
    if extension in {".html", ".htm"}:
        soup = BeautifulSoup(decoded, "lxml")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        lines = []
        for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
            text = " ".join(element.get_text(" ", strip=True).split())
            if not text:
                continue
            lines.append(f"## {text}" if element.name in {"h1", "h2", "h3"} else text)
        return "\n\n".join(lines)
    return decoded


async def load_local_document(
    *,
    filename: str,
    content: bytes,
    source_url: str,
    language: str = "zh-CN",
    country_or_region: str | None = None,
    business_category: str | None = None,
) -> ParsedDocument:
    """Parse an uploaded document without blocking the event loop."""
    text = await asyncio.to_thread(_load_sync, filename, content)
    if not text.strip():
        raise AppError("文件未提取到有效文本", code="empty_document", status_code=422)
    return ParsedDocument(
        title=PurePath(filename).stem,
        source_url=source_url,
        source_type=PurePath(filename).suffix.lower().lstrip("."),
        content=text,
        language=language,
        country_or_region=country_or_region,
        business_category=business_category,
        crawled_at=datetime.now(UTC),
    )
